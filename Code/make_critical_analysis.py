"""Generates Paper_Critical_Analysis.docx -- a structured critique of the selected
base paper (Jouffe, 1998) demonstrating understanding of its methodology, strengths,
and limitations, ahead of the proposed extension. Headline numbers are read live
from Results/ so this document cannot drift from the manuscript.
"""
import json
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = Path(__file__).resolve().parent
ROOT = HERE.parent
RESULTS = ROOT / "Results"
OUT = ROOT / "Supplementary" / "Paper_Critical_Analysis.docx"


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build():
    with open(RESULTS / "metrics.json") as f:
        metrics = json.load(f)
    with open(RESULTS / "statistical_tests.json") as f:
        stats_j = json.load(f)
    base = metrics["Baseline FQL"]
    gta = metrics["Improved FQL-GTA"]
    n_seeds = base["n_seeds_total"]
    pc = stats_j["primary_comparison"]
    return_gain_pct = (gta["final_return_mean"] - base["final_return_mean"]) / base["final_return_mean"] * 100
    auc_gain_pct = (gta["auc_mean"] - base["auc_mean"]) / base["auc_mean"] * 100

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Critical Analysis of the Selected Research Paper", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Fuzzy Inference System Learning by Reinforcement Methods (Jouffe, 1998)")
    run.italic = True
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2.add_run("Karthikeyan K  |  MS in Artificial Intelligence & Machine Learning, REVA University  |  "
                 "Advisor: Dr. J. B. Simha")

    add_heading(doc, "1. Paper Identification", level=1)
    add_body(doc,
        "Jouffe, L. (1998). Fuzzy inference system learning by reinforcement methods. "
        "IEEE Transactions on Systems, Man, and Cybernetics -- Part C: Applications and "
        "Reviews, 28(3), 338-355. DOI: 10.1109/5326.704563.")

    add_heading(doc, "2. Problem the Paper Addresses", level=1)
    add_body(doc,
        "Classical reinforcement learning methods such as Q-learning assume a discrete, "
        "enumerable set of states. Real control problems are continuous, so a tabular "
        "value function is either impossible or prohibitively large. The paper addresses "
        "how to learn a control policy by reinforcement when the state space is "
        "continuous, while keeping the learned controller interpretable. The proposed "
        "answer is to use a fuzzy inference system as the structure that both discretises "
        "the state softly and stores the learned action values.")

    add_heading(doc, "3. Methodology of the Original Work (as Understood)", level=1)
    add_body(doc,
        "The state space is partitioned by fuzzy sets defined along each input dimension. "
        "Combining the sets across dimensions produces a rule base; each rule carries a "
        "set of competing candidate actions together with a quality value for each. Given "
        "an observed state, the degree to which each rule is active (its firing strength) "
        "is computed by a t-norm across the membership degrees of the participating fuzzy "
        "sets. Rather than selecting a single active state, the method blends the active "
        "rules: an action is chosen per rule, and the global action is inferred from the "
        "firing-strength-weighted contributions of the rules.")
    add_body(doc,
        "Learning proceeds by temporal-difference reinforcement. After a transition, the "
        "method forms an estimate of the value of the current situation and compares it "
        "with the reward obtained plus the discounted value of the next situation. The "
        "resulting temporal-difference error is used to adjust the quality values of the "
        "rules that were responsible for the chosen action, in proportion to how strongly "
        "each rule fired. Exploration is handled by occasionally trying non-greedy "
        "candidate actions so that the quality estimates are not trapped in an early local "
        "preference. The pole-balancing (cart-pole) task is used as a representative "
        "continuous-control demonstration.")

    add_heading(doc, "4. Strengths of the Original Work", level=1)
    add_bullets(doc, [
        "Interpretability. The learned controller remains a set of linguistic fuzzy rules "
        "that a domain engineer can read and audit -- a decisive advantage over opaque "
        "function approximators.",
        "Continuous-state capability. The fuzzy partition allows reinforcement learning to "
        "operate on continuous inputs without a hand-crafted discrete grid, while the soft "
        "membership gives graceful generalisation between neighbouring situations.",
        "Simplicity and low cost. The update is a lightweight linear adjustment of rule "
        "parameters, with no gradient back-propagation through a deep network, making it "
        "fast and stable on modest hardware.",
    ])

    add_heading(doc, "5. Limitations Identified (the Basis for the Extension)", level=1)
    add_body(doc,
        "Three limitations of the original formulation are the specific focus of the "
        "extension proposed in this submission (see "
        "Manuscript/FQL-GTA_Final_Manuscript.docx for the full treatment):")
    add_bullets(doc, [
        "Triangular membership functions are piecewise-linear and non-differentiable at "
        "their vertices; adjacent rules hand off activation abruptly rather than blending "
        "smoothly, which coarsens the effective resolution of the state representation "
        "near rule boundaries.",
        "The one-step temporal-difference update credits only the single most recent "
        "state-action pair on every step. A reward realised several steps after the "
        "decision that caused it propagates back one step at a time, which is slow when "
        "the credit-assignment path is long, as it often is in balancing-type control "
        "tasks.",
        "A fixed exploration rate held constant for the whole run over-explores late in "
        "training, once the agent already has a reasonable policy, and under-explores "
        "early, when the rule base is still essentially uninformed.",
    ])

    add_heading(doc, "6. Proposed Extension and How It Addresses Each Limitation", level=1)
    add_body(doc,
        "The extension, termed FQL-GTA, replaces each of the three elements above with a "
        "principled alternative -- Gaussian membership functions, Watkins' Q(lambda) "
        "eligibility traces, and an adaptive epsilon schedule -- while holding rule "
        "resolution and the discount factor identical. The two agents' learning rates "
        "differ by frozen configuration (baseline 0.3, FQL-GTA 0.2); a dedicated "
        "sensitivity sweep (manuscript Section 6.1) confirms this does not drive the "
        "primary comparison's result.")
    add_body(doc,
        f"The comparative study confirms the expected effects: on CartPole-v1 across "
        f"{n_seeds} independent seeds, the extension raised the mean final trailing-window "
        f"return by {return_gain_pct:.1f}% ({base['final_return_mean']:.1f} to "
        f"{gta['final_return_mean']:.1f}), improved area-under-the-learning-curve by "
        f"{auc_gain_pct:.1f}%, and lifted the strict solve rate (trailing-50 mean still "
        f"above threshold at episode 500) from {base['solved_rate']:.0%} to "
        f"{gta['solved_rate']:.0%}, confirmed by a paired Wilcoxon signed-rank test "
        f"(p = {pc['wilcoxon_p_value']:.2e}, Cohen's d = "
        f"{pc['cohens_d_trailing50_return']:.2f}), while the interpretable fuzzy rule "
        f"base was preserved. A full 2-cubed factorial ablation further shows no single "
        f"component reaches the solve threshold alone, with significant two-way "
        f"interactions between Gaussian membership functions and eligibility traces and "
        f"between Gaussian membership functions and adaptive exploration, with no "
        f"significant three-way interaction -- so the gain is not attributable to any "
        f"single component.")

    add_heading(doc, "7. Relation to Recent Literature", level=1)
    add_body(doc,
        "The chosen direction is consistent with, but distinct from, recent fuzzy-RL "
        "research. Neuro-fuzzy controllers have been trained with deep value methods and, "
        "more recently, with policy-gradient algorithms; interpretable-control studies "
        "have generated compact fuzzy controllers from batch data; and fuzzy-evaluated "
        "reinforcement models have been applied to fault forecasting for maintenance. "
        "Whereas those lines replace or wrap the fuzzy system with heavier machinery, the "
        "present extension strengthens the classical tabular-fuzzy agent itself with "
        "three lightweight modifications, keeping it interpretable and inexpensive -- an "
        "appropriate and defensible scope for a focused reproduction-and-improvement "
        "study, evaluated here with the statistical rigour (30 seeds, full factorial "
        "ablation, formal cross-environment transfer study) a journal submission "
        "requires.")

    add_heading(doc, "8. Conclusion of the Analysis", level=1)
    add_body(doc,
        "The selected paper is a sound, foundational contribution whose methodology is "
        "well understood and faithfully reproduced in this work. Its three principal "
        "limitations are clearly identifiable and individually addressable, which makes "
        "it an excellent basis for a meaningful, novel-yet-tractable extension. The "
        "experimental comparison demonstrates that the proposed enhancements deliver "
        "substantial, statistically validated gains without sacrificing the "
        "interpretability that motivates fuzzy reinforcement learning in the first "
        "place -- while a full factorial ablation and a formal cross-environment "
        "transfer study also report, honestly, where the evidence is weaker (the "
        "ablation's pairwise comparisons are underpowered at 3 seeds per condition; "
        "zero-shot hyperparameter transfer to new environments is weak).")

    doc.save(str(OUT))
    print(f"Saved {OUT}")


if __name__ == "__main__":
    build()
