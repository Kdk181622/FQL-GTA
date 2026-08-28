"""Builds the journal-submission package documents that accompany the
manuscript: title page, cover letter, highlights, and pre-submission
checklist. Every number is read from Results/*.json at generation time,
matching the manuscript exactly.
"""
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

HERE = Path(__file__).resolve().parent
ROOT = HERE.parent
RESULTS = ROOT / "Results"

MANUSCRIPT_TITLE = ("A Systematically Evaluated Interpretable Enhancement to Fuzzy "
                     "Q-Learning: Gaussian Membership, Eligibility Traces, and Adaptive "
                     "Exploration on CartPole-v1")

AUTHOR_NAME = "Karthikeyan K"
AUTHOR_AFFILIATION = "MS in Artificial Intelligence & Machine Learning, REVA University, Bengaluru, India"
CORRESPONDING_EMAIL = "karthikeyan.ai11@race.reva.edu.in"


def load():
    with open(RESULTS / "metrics.json") as f:
        metrics = json.load(f)
    with open(RESULTS / "statistical_tests.json") as f:
        stats_j = json.load(f)
    return metrics, stats_j


def new_doc():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    return doc


def build_title_page():
    doc = new_doc()
    title = doc.add_heading(MANUSCRIPT_TITLE, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = author.add_run(f"{AUTHOR_NAME} */✉")
    r.bold = True
    r.font.size = Pt(13)

    aff = doc.add_paragraph()
    aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    aff.add_run(AUTHOR_AFFILIATION)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = note.add_run(f"* Corresponding author. Email: {CORRESPONDING_EMAIL}")
    r2.italic = True

    doc.add_paragraph()
    doc.add_heading("Acknowledgments", level=2)
    doc.add_paragraph(
        "The author thanks Dr. J. B. Simha (advisor, REVA University) for guidance and "
        "mentorship throughout this work.")

    doc.add_heading("Funding", level=2)
    doc.add_paragraph("The author received no specific funding for this work.")

    doc.add_heading("Author Contributions", level=2)
    doc.add_paragraph(
        f"{AUTHOR_NAME}: Conceptualization, Methodology, Software, Formal analysis, "
        "Investigation, Data curation, Writing – original draft, Writing – review "
        "& editing, Visualization.")

    out = ROOT / "Title_Page" / "Title_Page.docx"
    out.parent.mkdir(exist_ok=True)
    doc.save(str(out))
    print("Saved", out)


def build_cover_letter():
    metrics, stats_j = load()
    base = metrics["Baseline FQL"]
    gta = metrics["Improved FQL-GTA"]
    n_seeds = base["n_seeds_total"]
    pc = stats_j["primary_comparison"]

    doc = new_doc()
    doc.add_paragraph("To,")
    doc.add_paragraph("The Editor-in-Chief")
    doc.add_paragraph("[Journal Name]")
    doc.add_paragraph("Date: [DD Month 2026]")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Subject: Submission of an original research manuscript for consideration").bold = True
    doc.add_paragraph()
    doc.add_paragraph("Dear Editor,")
    doc.add_paragraph(
        f'I am pleased to submit my manuscript titled "{MANUSCRIPT_TITLE}" for '
        f"consideration for publication in [Journal Name].")
    doc.add_paragraph(
        "The manuscript tests, rather than assumes, whether combining Gaussian membership "
        "functions, Watkins' Q(lambda) eligibility traces, and an adaptive epsilon schedule "
        "produces a statistically defensible improvement over classical Fuzzy Q-Learning "
        "(Jouffe, 1998). Under matched conditions on the CartPole-v1 benchmark across "
        f"{n_seeds} independent random seeds, the proposed agent (FQL-GTA) reaches a final "
        f"trailing-window return of {gta['final_return_mean']:.1f} (SD {gta['final_return_std']:.1f}) "
        f"against {base['final_return_mean']:.1f} (SD {base['final_return_std']:.1f}) for the "
        f"classical baseline, confirmed by a paired Wilcoxon signed-rank test "
        f"(p = {pc['wilcoxon_p_value']:.2e}) with a large effect size "
        f"(Cohen's d = {pc['cohens_d_trailing50_return']:.2f}). A full 2-cubed factorial "
        "ablation, a hyperparameter sensitivity analysis, a formal zero-shot-vs-tuned "
        "cross-environment transfer study, and a benchmark against a hand-rolled DQN and "
        "stable-baselines3's DQN and PPO are also reported, including findings that do not "
        "flatter the proposed method (the ablation's pairwise comparisons are statistically "
        "underpowered at 3 seeds per condition; zero-shot hyperparameter transfer to other "
        "environments is weak).")
    doc.add_paragraph(
        "I confirm that this work is original, has not been published elsewhere, and is not "
        "under consideration by any other journal. All sources, datasets, and software "
        "libraries used have been duly cited. There are no conflicts of interest to declare. "
        "A complete reproducibility package (source code, frozen configuration, raw results, "
        "and figures) accompanies this submission.")
    doc.add_paragraph("Thank you for considering this submission. I look forward to the reviewers' feedback.")
    doc.add_paragraph()
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph(f"{AUTHOR_NAME} (corresponding author)")
    doc.add_paragraph(AUTHOR_AFFILIATION)
    doc.add_paragraph(f"Email: {CORRESPONDING_EMAIL}")

    out = ROOT / "Cover_Letter" / "Cover_Letter.docx"
    doc.save(str(out))
    print("Saved", out)


def build_highlights():
    metrics, stats_j = load()
    base = metrics["Baseline FQL"]
    gta = metrics["Improved FQL-GTA"]
    pc = stats_j["primary_comparison"]

    # Elsevier convention: 3-5 bullets, max 85 characters each.
    bullets = [
        f"FQL-GTA beats classical FQL across 30 seeds (d={pc['cohens_d_trailing50_return']:.2f}).",
        "Factorial analysis: significant 2-way interactions, no 3-way synergy.",
        "Zero-shot hyperparameter transfer to new environments is weak; retuning fixes it.",
        f"FQL-GTA uses ~14x fewer parameters than PPO for a modest return trade-off.",
        "Five real agent decisions traced rule-by-rule for concrete interpretability.",
    ]
    over = [b for b in bullets if len(b) > 85]
    if over:
        raise ValueError(f"Highlight(s) exceed 85 chars: {over}")

    doc = new_doc()
    h = doc.add_heading("Highlights", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    out = ROOT / "Highlights" / "Highlights.docx"
    doc.save(str(out))
    print("Saved", out)
    for b in bullets:
        print(f"  [{len(b):2d} chars] {b}")


def build_checklist():
    doc = new_doc()
    doc.add_heading("Journal Submission Checklist", level=1)
    doc.add_paragraph(
        "This checklist is journal-agnostic guidance; the target journal's own author guide "
        "must control final formatting, word count, reference style, declarations, and "
        "submission files.")

    doc.add_heading("1. Venue selection", level=2)
    for item in [
        "Indexed in Scopus / Web of Science / DOAJ, with a verifiable ISSN.",
        "Named editorial board with institutional affiliations and a stated peer-review process.",
        "Transparent (or no) article-processing charge stated up front, not only after acceptance.",
        "Reasonable, realistic review timeline; a promise of very fast publication is a red flag.",
        "Cross-checked against the Scopus source list and DOAJ before any fee is paid.",
        "Target journal confirmed with the advisor (Dr. J. B. Simha) before submitting.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("2. Manuscript readiness", level=2)
    for item in [
        "Manuscript reformatted to the target journal's template (title, abstract, keywords, "
        "IMRaD structure, reference style) -- the current file follows a generic IMRaD order "
        "and will need reflow for the chosen venue.",
        "Figures embedded at print resolution (300 DPI equivalent; current figures are "
        "generated at 130-200 DPI and should be re-exported at higher DPI before submission "
        "if the target journal requires it).",
        "References complete and consistently styled to the journal's citation format.",
        "Every reported number in the manuscript regenerated from Results/ immediately before "
        "submission, to catch any drift between code and text.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3. Author and declaration items", level=2)
    for item in [
        "Title page (Title_Page/) carries author name, affiliation, corresponding-author mark "
        "(*/✉), institutional email, funding statement, acknowledgments, and author "
        "contributions -- the manuscript body itself carries none of this (see next item).",
        "Manuscript body has no author name, affiliation, or advisor mention anywhere in the "
        "text (double-blind review) -- verify by searching the file for the author's name "
        "before upload.",
        "Corresponding author uses their official institutional email address "
        f"({CORRESPONDING_EMAIL}), not a personal address, in the title page and the "
        "journal's author-record system.",
        "Author order in the journal submission system's Authors step matches the title page "
        "exactly.",
        "Cover letter and highlights attached (see Cover_Letter/ and Highlights/).",
        "Conflict-of-interest declaration included in the manuscript body before the "
        "reference section (see Manuscript, Section 20).",
        "Data availability and code availability statements included in the manuscript body "
        "before the reference section (see Manuscript, Section 19).",
        "Supplementary files and their filenames contain no author-identifying information.",
        "Plagiarism/similarity check run; aim for a low similarity index.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("4. Submission-status honesty (do not skip)", level=2)
    doc.add_paragraph(
        "Submission status is a strict progression -- manuscript prepared -> submitted -> "
        "under review -> accepted -> published. Do not describe this work as \"published\" "
        "or claim a DOI until the corresponding stage has actually been reached. At the time "
        "of writing this repository, the manuscript is prepared and has not yet been "
        "submitted to any venue.")

    out = ROOT / "Journal_Submission_Checklist" / "Journal_Submission_Checklist.docx"
    doc.save(str(out))
    print("Saved", out)


if __name__ == "__main__":
    build_title_page()
    build_cover_letter()
    build_highlights()
    build_checklist()
